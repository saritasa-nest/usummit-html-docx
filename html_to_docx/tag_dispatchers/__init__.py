import re

from docx.document import Document
from docx.enum.text import WD_UNDERLINE
from docx.table import _Cell
from docx.text.paragraph import Paragraph


class ParentTagMixin:
    """Applies formatting from parent elements."""

    @classmethod
    def _apply_parent_formatting(cls, element, run):
        """Applies formatting implemented in a parent element."""
        if not element.getparent():
            return run
        if element.getparent().tag in ['em', 'i']:
            run.italic = True
        if element.getparent().tag in ['strong', 'b']:
            run.bold = True
        if element.getparent().tag == 'u':
            run.underline = WD_UNDERLINE.SINGLE

        # Applies formatting from parent's parent. Example:
        # <u><strong><em>text</em></strong></u>
        run = cls._apply_parent_formatting(element.getparent(), run)
        return run


class ParagraphTailMixin:

    @classmethod
    def append_tail(cls, element, container):
        """Appends tail as a new paragraph."""
        paragraph = cls.get_new_paragraph(container)
        text = replace_whitespaces(element.tail).strip()
        if not text:
            return container

        style = None
        if element.getparent().tag == 'blockquote':
            style = 'Intense Quote Char'

        paragraph.add_run(text=text, style=style)
        return container


class CharacterTailMixin:

    @classmethod
    def append_tail(cls, element, container):
        """Appends tail to the current paragraph."""
        text = replace_whitespaces(element.tail)
        run = container.add_run(text=text)
        cls._apply_parent_formatting(element, run)
        return container


class TagDispatcher(object):
    @classmethod
    def append_head(cls, element, container):
        raise NotImplementedError("Implemented in inheriting classes")

    @classmethod
    def append_tail(cls, element, container):
        raise NotImplementedError("Implemented in inheriting classes")

    @classmethod
    def get_current_paragraph(cls, container):
        current_paragraph = container
        if isinstance(container, Paragraph):
            if isinstance(container._parent, _Cell):
                current_paragraph = container._parent.paragraphs[-1]

        if isinstance(container, Document):
            current_paragraph = container.add_paragraph()
        return current_paragraph

    @classmethod
    def get_new_paragraph(cls, container):
        new_paragraph = container
        if isinstance(container, Paragraph):
            if isinstance(container._parent, _Cell):
                new_paragraph = container._parent.paragraphs[0]
                if len(container._parent.paragraphs) > 1:
                    new_paragraph = container._parent.add_paragraph()
                else:
                    if container._parent.paragraphs[0].text:
                        new_paragraph = container._parent.add_paragraph()
            else:
                if container.text:
                    new_paragraph = container._parent.add_paragraph()
        if isinstance(container, Document):
            new_paragraph = container.add_paragraph()
        return new_paragraph


def replace_whitespaces(text):
    """
    replaces multiple whitespaces and line breaks by a single whitespace
    """
    if text:
        text = ' '.join(text.split('\n'))
        text = re.sub(' +', ' ', text)

    return text if text else ''
