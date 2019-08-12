# encoding: utf-8
from html_docx.html.tag_dispatchers import TagDispatcher, replace_whitespaces
from base64 import b64decode
from io import BytesIO
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

MAX_WIDTH = Inches(7.5)
MAX_LENGTH = Inches(8)

class ImgDispatcher(TagDispatcher):
    def __init__(self):
        super(ImgDispatcher, self).__init__()

    @classmethod
    def append_head(cls, element, container):
        container = cls.get_current_paragraph(container)
        return cls._append_img(element, container)

    @classmethod
    def append_tail(cls, element, container):

        return cls._append_img(element, container)

    @classmethod
    def _append_img(cls, element, container):
        """
        First real add on to this thing -kevin
        <img> lets see if this can work.....
        """
        base64image = element.attrib['src'].split('base64,')[1]
        image_filelike =  BytesIO(b64decode(base64image))

        #container =container._parent.add_paragraph()
        run = container.add_run()
        run.add_break()
        inline_shape = run.add_picture(image_filelike)

        if inline_shape.width > MAX_WIDTH:
            scalar = MAX_WIDTH/inline_shape.width
            inline_shape.width = int(scalar* inline_shape.width)
            inline_shape.height = int(scalar* inline_shape.height)

        if inline_shape.height > MAX_LENGTH:
            scalar = MAX_LENGTH / inline_shape.height
            inline_shape.width = int(scalar * inline_shape.width)
            inline_shape.height = int(scalar * inline_shape.height)

        container.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        print("here")
       #container.add_picture(image_filelike)


        return container